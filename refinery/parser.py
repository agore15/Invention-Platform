import os
import json
from docx import Document
from docx.oxml.ns import qn
import sys
from dataclasses import dataclass, asdict
from typing import List, Dict, Optional

# Add project root to sys.path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

@dataclass
class Chunk:
    text: str
    metadata: Dict[str, str]

class TDocParser:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None
        self.metadata = {}
        self.content = ""
        self.chunks: List[Chunk] = []
        self.cr_fields = {}

    def parse(self):
        """Main parsing method."""
        if not os.path.exists(self.file_path):
            print(f"Error: File not found {self.file_path}")
            return None

        try:
            self.doc = Document(self.file_path)
            self._extract_metadata()
            self._extract_cr_fields()
            self._extract_body_text_and_chunks()
            
            return {
                "metadata": self.metadata,
                "cr_fields": self.cr_fields,
                "content": self.content,
                "chunks": [asdict(c) for c in self.chunks]
            }
            
        except Exception as e:
            print(f"Error parsing {self.file_path}: {e}")
            return None

    def _extract_metadata(self):
        """Extracts metadata from the document header/first page."""
        # Check tables first (for CRs)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text.startswith("Source:"):
                        self.metadata["source"] = text.replace("Source:", "").strip()
                    elif "Source:" in text:
                        parts = text.split("Source:")
                        if len(parts) > 1:
                            self.metadata["source"] = parts[1].strip()
                    
                    if text.startswith("Title:"):
                        self.metadata["title"] = text.replace("Title:", "").strip()
                    elif "Title:" in text:
                        parts = text.split("Title:")
                        if len(parts) > 1:
                            self.metadata["title"] = parts[1].strip()

        # Check paragraphs (for LS) if missing
        if not self.metadata.get("source") or not self.metadata.get("title"):
            for para in self.doc.paragraphs[:30]: # Check first 30 paragraphs
                text = para.text.strip()
                if not self.metadata.get("source"):
                    if text.startswith("Source:"):
                        self.metadata["source"] = text.replace("Source:", "").strip()
                if not self.metadata.get("title"):
                    if text.startswith("Title:"):
                        self.metadata["title"] = text.replace("Title:", "").strip()

    def _extract_cr_fields(self):
        """Extracts Reason for Change and Summary of Change from CRs."""
        for table in self.doc.tables:
            for row in table.rows:
                for i, cell in enumerate(row.cells):
                    text = cell.text.strip().lower()
                    if "reason for change" in text:
                        if i + 1 < len(row.cells):
                            self.cr_fields["reason_for_change"] = row.cells[i+1].text.strip()
                    elif "summary of change" in text:
                        if i + 1 < len(row.cells):
                            self.cr_fields["summary_of_change"] = row.cells[i+1].text.strip()

    def _extract_body_text_and_chunks(self):
        """Extracts the main body text and creates chunks, handling Track Changes."""
        full_text = []
        current_section = "General"
        
        for para in self.doc.paragraphs:
            text = self._get_para_text_clean(para)
            if not text:
                continue
                
            # Simple heuristic for section headers
            if len(text) < 100 and (text[0].isdigit() or text.isupper()):
                 current_section = text
            
            full_text.append(text)
            
            if len(text) > 20: 
                chunk_metadata = self.metadata.copy()
                chunk_metadata["section"] = current_section
                self.chunks.append(Chunk(text=text, metadata=chunk_metadata))
        
        self.content = "\n".join(full_text)

    def _get_para_text_clean(self, paragraph):
        """Extracts text from a paragraph, ignoring deletions and including insertions."""
        text = ""
        try:
            p_element = paragraph._element
            for child in p_element:
                tag = child.tag
                if tag.endswith('r'): # Normal run
                    if child.text:
                        text += child.text
                elif tag.endswith('ins'): # Insertion
                    for sub in child:
                        if sub.tag.endswith('r'):
                            if sub.text:
                                text += sub.text
                elif tag.endswith('del'): # Deletion
                    pass
                elif tag.endswith('hyperlink'):
                    for sub in child:
                         if sub.tag.endswith('r'):
                            if sub.text:
                                text += sub.text
        except Exception:
            return paragraph.text.strip()
                            
        return text.strip()

if __name__ == "__main__":
    # Simple test if run directly
    pass
